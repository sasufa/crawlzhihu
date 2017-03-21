import requests
from bs4 import BeautifulSoup
import re
import json
import time
import http.cookiejar
from  PIL import Image
from docx import Document
import os
import random
import threadpool
import threading
from docx.shared import Inches

class Tool:
    RemoveImg = re.compile('<img.*?>| {7}|')
    # 匹配图片标签
    RemoveAddr = re.compile('<a.*?>|</a>')  # 匹配链接标签
    ReplaceLine = re.compile('<tr>|<div>|</div></p>')  # 匹配换行符的标签
    ReplaceTD = re.compile('<td>')  # 匹配制表符
    ReplacePara = re.compile('<p.*?>')  # 匹配段落开头
    ReplaceBR = re.compile('<br><br>|<br>')  # 匹配换行和双换行
    RemoveTag = re.compile('<.*?>')  # 匹配其余的标签

    def replace(self, x):
        x = re.sub(self.RemoveImg, "", x)  # 删除图片
        x = re.sub(self.RemoveAddr, "", x)  # 删除链接
        x = re.sub(self.ReplaceLine, "\n", x)  # 替换换行符
        x = re.sub(self.ReplaceTD, "\t", x)  # 将制表符替换
        x = re.sub(self.ReplacePara, "\n  ", x)  # 段落开头替换
        x = re.sub(self.ReplaceBR, "\n", x)
        x = re.sub(self.RemoveTag, "", x)  # 删除其余标签
        return x.strip()




class Zhihu():
    def __init__(self, url):
        self.url = url
        self.start()
        #self.file = None
        #self.file2 = None
    def loginzhihu(self,):
        def get_xsrf():
            """
            获取参数_xsrf
            """
            response = s.get('https://www.zhihu.com', headers=headers)
            html = response.text
            get_xsrf_pattern = re.compile(r'<input type="hidden" name="_xsrf" value="(.*?)"')
            _xsrf = re.findall(get_xsrf_pattern, html)[0]
            return _xsrf

        def get_captcha():
            """
            获取验证码本地显示
            返回你输入的验证码
            """
            t = str(int(time.time() * 1000))
            captcha_url = 'http://www.zhihu.com/captcha.gif?r=' + t + "&type=login"
            response = s.get(captcha_url, headers=headers)
            with open('zhihucptcha.gif', 'wb') as f:
                f.write(response.content)
            im = Image.open('zhihucptcha.gif')
            im.show()
            captcha = input('本次登录需要输入验证码： ')
            return captcha

        def login(username, password):
            # 检测到11位数字则是手机登录
            if re.match(r'\d{11}$', account):
                print('使用手机登录中...')
                loginurl = 'http://www.zhihu.com/login/phone_num'
                data = {'_xsrf': get_xsrf(),
                        'password': password,
                        'remember_me': 'true',
                        'phone_num': username
                        }
            else:
                print('使用邮箱登录中...')
                loginurl = 'https://www.zhihu.com/login/email'
                data = {'_xsrf': get_xsrf(),
                        'password': password,
                        'remember_me': 'true',
                        'email': username
                        }
            # 若不用验证码，直接登录

            result = s.post(loginurl, data=data, headers=headers)
            msg = (json.loads(result.text))['msg']
            print(msg)
            # 要用验证码，post后登录
            if not "成功" in msg:
                data['captcha'] = get_captcha()
                result = s.post(loginurl, data=data, headers=headers)
                print((json.loads(result.text))['msg'])
            # 保存cookie到本地
            s.cookies.save(ignore_discard=True, ignore_expires=True)
        if __name__ == '__main__':
            account = "你的知乎账户"
            secret = "你的知乎密码"
            login(account, secret)
            # 设置里面的简介页面，登录后才能查看。以此来验证确实登录成功
            get_url = 'https://www.zhihu.com/settings/profile'
            # allow_redirects=False 禁止重定向
            resp = s.get(get_url, headers=headers, allow_redirects=False)
            print(resp.text)

    # 获取网页代码
    def getsoup(self, url):
        #s = requests.session()
        self.data = s.get(url, headers=headers)
        soup = BeautifulSoup(self.data.content, 'lxml')
        return soup

        # 获取帖子标题
    def gettitle(self, soup):
        title = soup.title.string.replace("\n", "")  # 问题题目
        title = title.replace(r"/", "")
        return title

    def filename(self, title):
         if title:
             self.doc = Document()
             #self.file = open(title + ".txt", "w+", encoding="utf-8")
             #self.file2 = open(title + ".txt", "w+", encoding="utf-8")
         else:
             print("文件创建出错！")

            # 获取题目描述
    def getdescription(self, soup):
        description = soup.find('div', {'class': 'zm-editable-content'}).strings  # 问题描述，可能多行
        detail = ""
        for each in description:
            detail += each
        detail = detail.replace("\n", "")
        self.doc.add_paragraph("\n问题描述为:\n" + detail + '\n')
        return (detail)

    def getpic (self, imgurl, file, name):
        print('downloading picture')
        try:
            file = "知乎图片\\" + file
            if not os.path.exists(file):
                os.makedirs(file)
            path = file + "\\" + name
            ir = requests.get(imgurl)
            if ir.status_code == 200 and not os.path.exists(path) :
                open(path, 'wb').write(ir.content)
        except AttributeError as e:
            path = r'C:\Users\Lisen\Desktop\error.jpg'
            print('download picture error, the link is ' + imgurl)
            pass
        finally:
            return path

  # 获取回答
    def getanswer(self, soup, title):
        # 保存网页代码
        #self.file2.write(str(soup.contents))
        #baocunwangyedaima
        #s = requests.session()
        answer_num = int(soup.find('h3', {'id': 'zh-question-answer-num'}).string.split(' ')[0])  # 答案数量
        self.doc.add_paragraph("\n本问题共" + str(answer_num) + '个回答\n')
        #self.file.write("\n本问题共" + str(answer_num) + '个回答\n')
        num = 1
        index = soup.find_all('div', {'tabindex': '-1'})
        for i in range(len(index)):
            print('正在抓取第 ' + str(num) + '个答案，共' + str(answer_num) + '个答案……' )
            if num == 2:
                print("1")
            #获取回答人信息
            try:
                a = index[i].find('a', {'class': 'author-link'})
                auth = str(num) + '__' + a.string
                href = 'http://www.zhihu.com' + a['href']
                self.doc.add_paragraph("\n" + auth + '\n' + href + "\n")
                #self.file.write("\n"+ auth + '\n'+ href + "\n")
            except:
                auth = str(num) + '__匿名用户'
                self.doc.add_paragraph("\n" + auth + '\n')
               # self.file.write("\n" + auth + '\n')
            #获取该答案点赞数
            try:
                vote = index[i].find('span', {'class': "count"})
                self.doc.add_paragraph("\n赞同数为：" + vote.string  + "\n")
                #self.file.write("\n赞同数为：" + vote.string  + "\n")
            except:
                self.doc.add_paragraph("\n赞同数为：  ---\n")
                #self.file.write("\n赞同数为：  ---\n")
            #获得答案内容
            try:
                answer_content = index[i].find('div', {'class': 'zm-editable-content clearfix'}).contents
            except:
                answer_content = ['作者修改内容通过后，回答会重新显示。如果一周内未得到有效修改，回答会自动折叠。']
            for content in answer_content:
                imgurl = None
                if str(type(content)) == "<class 'bs4.element.NavigableString'>" :
                    self.doc.add_paragraph(content + '\n')
                elif content.find('img', {'class': 'content_image'}):
                        # 获得答案中的图片
                    try:
                        content = content.find('img', {'class': 'content_image'})
                        imgurl = content['src']
                    except:
                        imgurl = None
                elif content.find('img', {'class': 'origin_image zh-lightbox-thumb'}):
                    try:
                        content = content.find('img', {'class': 'origin_image zh-lightbox-thumb'})
                        imgurl = content['src']
                    except:
                        imgurl = None
                if imgurl:
                    path = self.getpic(imgurl, title, imgurl[-10:])
                    try:
                        self.doc.add_picture(path,width=Inches(6))
                        #self.file.write(content + '\n')
                    except:
                        self.doc.add_paragraph("图片加载错误\nurl为" + imgurl)
            num += 1

        xsrf = soup.find('input', {'name': '_xsrf', 'type': 'hidden'}).get('value')
        url_token = re.findall('url_token(.*)', self.data.content.decode('utf-8'))[0][8:16]
        print(url_token)


        # 循环次数
        if answer_num % 10 == 0 :
            n = int(answer_num / 10)
        else:
            n = int (answer_num / 10 )+ 1
        #self.file3 = open("1212" + ".html", "w+", encoding="utf-8")
        for j in range(1, n + 1, 1):
            offset = 10 * j
            params = json.dumps({"url_token": url_token, "pagesize": 10, "offset": offset})
            payload = {"method": "next", "params": params, "_xsrf": xsrf}
            click_url = 'https://www.zhihu.com/node/QuestionAnswerListV2'
            count = 1
            while True:
                try:
                    time.sleep(random.random())
                    data = s.post(click_url, data=payload, headers=headers)
                    data = json.loads(data.content.decode('utf-8'))
                    break
                except:
                    self.doc.save(title + ".docx")
                    count = count + 1
                    print("链接中断，文件已保存，等待20秒后重试")
                    time.sleep(20)
                    if count == 3 :
                        self.loginzhihu()
                        count = 1

            #保存网页代码
            #self.file2.write(str(data['msg']))

            for answer in data['msg']:
                print('正在抓取第 ' + str(num) + '个答案，共' + str(answer_num) +'个答案……' )
                soup1 = BeautifulSoup(answer,'lxml')
                # 获取回答人信息
                try:
                    a = soup1.find('a', {'class': 'author-link'})
                    auth = str(num) + '__' + a.string
                    href = 'http://www.zhihu.com' + a['href']
                    self.doc.add_paragraph("\n" + auth + '\n' + href + "\n")
                    #self.file.write("\n" + auth + '\n' + href + "\n")
                except:
                    auth= str(num) + '__匿名用户'
                    self.doc.add_paragraph("\n" + auth + '\n')
                   #self.file.write("\n" + auth + '\n')
                # 获取该答案点赞数
                try:
                    vote = soup1.find('span', {'class': "count"})
                    self.doc.add_paragraph("\n赞同数为：" + vote.string + "\n")
                    #self.file.write("\n赞同数为：" + vote.string + "\n")
                except:
                    self.doc.add_paragraph("\n赞同数为：  ---\n")
                    #self.file.write("\n赞同数为：  ---\n")
                # 获得答案内容
                try:
                    answer_content = soup1.find('div', {'class': 'zm-editable-content clearfix'}).contents
                except:
                    answer_content = ['作者修改内容通过后，回答会重新显示。如果一周内未得到有效修改，回答会自动折叠。']
                for content in answer_content:
                    imgurl = None
                    if str(type(content)) == "<class 'bs4.element.NavigableString'>":
                        self.doc.add_paragraph(content + '\n')
                    elif content.find('img', {'class': 'content_image'}):
                        # 获得答案中的图片
                        try:
                            content = content.find('img', {'class': 'content_image'})
                            imgurl = content['src']
                        except:
                            imgurl = None
                    elif content.find('img', {'class': 'origin_image inline-img zh-lightbox-thumb'}):
                        try:
                            content = content.find('img', {'class': 'content_image'})
                            imgurl = content['src']
                        except:
                            imgurl = None
                    if imgurl:
                        path = self.getpic(imgurl, title, imgurl[-10:])
                        try:
                            self.doc.add_picture(path,width=Inches(6))
                            # self.file.write(content + '\n')
                        except:
                            self.doc.add_paragraph("图片加载错误\nurl为" + imgurl)
                num += 1

    def start(self):
        filename = 'cookie'
        # 建立一个会话，可以把同一用户的不同请求联系起来；直到会话结束都会自动处理cookies
        global s
        s = requests.Session()
        # 建立LWPCookieJar实例，可以存Set-Cookie3类型的文件。
        # 而MozillaCookieJar类是存为'/.txt'格式的文件
        s.cookies = http.cookiejar.LWPCookieJar(filename)
        # 若本地有cookie则不用再post数据了
        try:
            s.cookies.load(filename=filename, ignore_discard=True)
        except:
            print('Cookie未加载！')
            print("重新登录知乎")
            self.loginzhihu()
        while True:
            try:
                soup = self.getsoup(self.url)
                print("已获得源代码，正在处理......")
                title = self.gettitle(soup)
                print("本次获取的问题是：" + title+"；问题编号为：" + self.url[-8:])
                break
            except:
                print("重新登录知乎")
                self.loginzhihu()
        print("正在创建存储文件......")
        self.filename(title)
        detail = self.getdescription(soup)
        print("问题描述为：\n" + detail + "\n")
        self.getanswer(soup, title)
        self.doc.save(title + self.url[-8:]+".docx")
        print("问题：\n"+ title +"\n结束")




class Favorite():
    def __init__(self, favoriteurl):
        self.url = favoriteurl
        self.questionlist = []

    def getsoup(self, url):
        s = requests.session()
        data = s.get(url, headers=headers)
        soup = BeautifulSoup(data.content, 'lxml')
        return soup

    def findmaxpage(self, soup):
        maxpage = soup.find('span', text="下一页").previous_element.previous_element
        return int(maxpage)
    def getquestions(self, url):
        while True:
            try:
                soup = self.getsoup(url)
                break
            except:
                print("链接中断，文件已保存，等待20秒后重试,位置1")
                time.sleep(20)
        questions = soup.find_all('a', href=re.compile('/question/\d{8}$'), target="_blank")
        #questionlist = []
        if self.counter_lock.acquire():
            for question in questions:
                print(question)
                self.questionlist.append(question['href'].split('/')[-1])
            print("页面"+ url[-2:] + "已结束")
            self.counter_lock.release()
        return self.questionlist
    def getallquestions(self):
        self.counter_lock = threading.Lock()
        soup = self.getsoup(self.url)
        maxpage = self.findmaxpage(soup)
        task_threads = []  # 存储线程
        urls = []
        for i in range(1,(maxpage + 1),1):
            url = 'https://www.zhihu.com/collection/20359746?page=' + str(i)
            urls.append(url)
        pool = threadpool.ThreadPool(3)
        requests = threadpool.makeRequests(self.getquestions,urls,)
        [pool.putRequest(req) for req in requests]
        pool.wait()
        return self.questionlist

    # url = 'https://www.zhihu.com/question/31784351'
    # headers = ('User-Agent','Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2785.104 Safari/537.36 Core/1.53.2141.400 QQBrowser/9.5.10219.400')
    # opener = request.build_opener()
    # opener.addheaders = [headers]
    # question = zhihu(url)
    # question.start()

#url = 'https://www.zhihu.com/question/40950470'
headers = {
    "User-Agent": 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/53.0.2785.104 Safari/537.36 Core/1.53.2141.400 QQBrowser/9.5.10219.400',
    "Referer": "http://www.zhihu.com/",
    'Host': 'www.zhihu.com',
}

def downzhihu():
    url = 'https://www.zhihu.com/question/23147606'
    q = Zhihu(url)

def getfav():
    favoriteurl = 'https://www.zhihu.com/collection/20359746?page=1'
    f = Favorite(favoriteurl)
    questionlist = f.getallquestions()
    File = open("问题列表11.txt","w+")
    for que in questionlist:
       File.write(que + "\n")
    File.close()

def main():
    questionlist = open("问题列表.txt","r").read().split("\n")
    urls = []
    for que in questionlist:
        url = 'https://www.zhihu.com/question/' + que
        urls.append(url)
    pool = threadpool.ThreadPool(3)
    requests = threadpool.makeRequests(Zhihu, urls, )
    [pool.putRequest(req) for req in requests]
    pool.wait()
    print("线程结束")

if __name__ == '__main__':
    #main()
    downzhihu()
    #getfav()

